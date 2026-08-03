from typing import List, Dict, Optional
import re
import json
import hashlib
import shelve
import os
import glob
import tempfile
import pandas as pd
import gradio as gr

# ─────────────────────────────────────────────
# Dependencias opcionales
# ─────────────────────────────────────────────
try:
    from langdetect import detect as _langdetect
    HAS_LANGDETECT = True
except ImportError:
    HAS_LANGDETECT = False

try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as np
    HAS_MATPLOTLIB = True
except ImportError:
    HAS_MATPLOTLIB = False

try:
    from docx import Document
    from docx.shared import Inches
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pypdf
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

# ─────────────────────────────────────────────
# Categorías / keywords
# ─────────────────────────────────────────────
CATEGORIES_ES = {
    "pantalla":        ["pantalla", "amoled", "lcd", "pulgadas", "resolución", "hdr"],
    "cámara":          ["cámara", "camara", "mp", "foto", "fotografía", "night", "noche", "zoom", "sensor"],
    "batería":         ["batería", "mah", "carga", "autonomía", "duración", "inalámbrica"],
    "rendimiento":     ["ram", "procesador", "cpu", "snapdragon", "mediatek", "velocidad", "juego"],
    "almacenamiento":  ["almacenamiento", "gb", "rom", "memoria interna"],
    "conectividad":    ["5g", "4g", "wifi", "bluetooth", "nfc", "conexión"],
    "precio_calidad":  ["precio", "calidad-precio", "relación calidad-precio", "económico", "caro", "barato", "vale la pena"],
    "diseño":          ["diseño", "material", "aluminio", "cristal", "acabado", "ergonomía", "peso"],
    "software":        ["software", "actualizaciones", "android", "ios", "interfaz", "bloatware"],
    "audio":           ["audio", "altavoces", "sonido", "estéreo", "auriculares"],
    "durabilidad":     ["resistencia", "ip68", "ip67", "resistente", "duradero", "caídas", "sumergible"],
}
CATEGORIES_EN = {
    "screen":         ["screen", "inch", "amoled", "lcd", "resolution", "hdr"],
    "camera":         ["camera", "mp", "photo", "night", "zoom", "sensor"],
    "battery":        ["battery", "mah", "charge", "fast charging", "wireless"],
    "performance":    ["ram", "processor", "cpu", "snapdragon", "mediatek", "speed", "gaming"],
    "storage":        ["storage", "gb", "rom", "internal memory"],
    "connectivity":   ["5g", "4g", "wifi", "bluetooth", "nfc", "connection"],
    "price_quality":  ["price", "value for money", "affordable", "expensive", "cheap", "worth it"],
    "design":         ["design", "material", "aluminum", "glass", "finish", "ergonomics", "weight"],
    "software":       ["software", "updates", "android", "ios", "interface", "bloatware"],
    "audio":          ["audio", "speakers", "sound", "stereo", "headphones"],
    "durability":     ["durability", "ip68", "ip67", "resistant", "durable", "drops", "waterproof"],
}

NUMERIC_CATS = {
    "pantalla", "screen", "cámara", "camera", "batería", "battery",
    "rendimiento", "performance", "almacenamiento", "storage",
    "conectividad", "connectivity",
}

# Palabras de sentimiento
POSITIVE_ES = {"excelente", "increíble", "potente", "brillante", "rápido", "perfecto",
               "superior", "óptimo", "fluido", "nítido", "eficiente", "duradero"}
NEGATIVE_ES = {"mediocre", "lento", "pobre", "malo", "débil", "insuficiente",
               "deficiente", "limitado", "anticuado", "básico"}
POSITIVE_EN = {"excellent", "incredible", "powerful", "brilliant", "fast", "perfect",
               "superior", "optimal", "smooth", "crisp", "efficient", "lasting"}
NEGATIVE_EN = {"mediocre", "slow", "poor", "bad", "weak", "insufficient",
               "lacking", "limited", "outdated", "basic"}

# Negadores (para detectar frases como "no es rápido")
NEGATORS_ES = {"no", "sin", "tampoco", "nunca", "ni", "jamás", "apenas"}
NEGATORS_EN = {"no", "not", "without", "never", "neither", "nor", "hardly"}

# Patrones de especificaciones
SPEC_PATTERNS_ES = {
    "pantalla":       r'(\d+\.?\d*)\s*(?:pulgadas|")',
    "cámara":         r'(\d+)\s*(?:mp|megapíxeles)',
    "batería":        r'(\d+)\s*(?:mah)',
    "ram":            r'(\d+)\s*(?:gb)\s+(?:de\s+)?ram',
    "almacenamiento": r'(\d+)\s*(?:gb|tb)\s+(?:de\s+)?(?:almacenamiento|rom|memoria)',
}
SPEC_PATTERNS_EN = {
    "screen":   r'(\d+\.?\d*)\s*(?:inch|")',
    "camera":   r'(\d+)\s*(?:mp|megapixels)',
    "battery":  r'(\d+)\s*(?:mah)',
    "ram":      r'(\d+)\s*(?:gb)\s+ram',
    "storage":  r'(\d+)\s*(?:gb|tb)\s+(?:storage|rom|memory)',
}

TIE_MARGIN = 0.05          # margen relativo para considerar "empate" entre scores
MAX_TEXT_CHARS = 20000     # límite de longitud de texto por producto
CACHE_FILE = os.path.join(tempfile.gettempdir(), "comparador_cache")
MAX_PRODUCTS = 5


# ═════════════════════════════════════════════
# Utilidades
# ═════════════════════════════════════════════
def detect_language(texts: List[str]) -> str:
    if not HAS_LANGDETECT:
        return "es"
    try:
        sample = " ".join(t for t in texts if t.strip())[:500]
        lang = _langdetect(sample)
        return "es" if lang == "es" else "en"
    except Exception:
        return "es"

def clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip())

def split_sentences(text: str) -> List[str]:
    if not text:
        return []
    sents = re.split(r'(?<=[.!?,])\s+', text)
    return [s.strip().strip('"\'') for s in sents if s.strip()]

def tokenize(text: str) -> List[str]:
    return re.findall(r"\w+", text.lower())

def is_tie(a: float, b: float, margin: float = TIE_MARGIN) -> bool:
    """Compara dos scores con un margen relativo en vez de igualdad estricta de floats."""
    scale = max(abs(a), abs(b), 1.0)
    return abs(a - b) <= margin * scale

def detect_categories(sent: str, language: str, kws_dict: Optional[Dict[str, List[str]]] = None) -> List[str]:
    cats = []
    if kws_dict is None:
        kws_dict = CATEGORIES_ES if language == "es" else CATEGORIES_EN
    s = sent.lower()
    for cat, kws in kws_dict.items():
        for kw in kws:
            kw_low = kw.lower()
            # \b en todas las categorías para evitar falsos positivos (ej. "5g" dentro de "5gb")
            if re.search(rf'\b{re.escape(kw_low)}\b', s):
                cats.append(cat)
                break
    return cats

def extract_specs(text: str, language: str) -> Dict[str, List[float]]:
    """Devuelve TODAS las coincidencias por spec (no solo la primera)."""
    patterns = SPEC_PATTERNS_ES if language == "es" else SPEC_PATTERNS_EN
    specs: Dict[str, List[float]] = {}
    t = text.lower()
    for key, pattern in patterns.items():
        matches = re.findall(pattern, t)
        if matches:
            specs[key] = [float(m) for m in matches]
    return specs

def aggregate_specs(specs: Dict[str, List[float]]) -> Dict[str, float]:
    """Agrega una lista de valores de spec a un único valor representativo (el máximo)."""
    return {k: max(v) for k, v in specs.items()}

def sentiment_score(sentence: str, language: str) -> str:
    """Cuenta palabras positivas/negativas usando tokenización real y detectando negación."""
    tokens = tokenize(sentence)
    pos_w = POSITIVE_ES if language == "es" else POSITIVE_EN
    neg_w = NEGATIVE_ES if language == "es" else NEGATIVE_EN
    negators = NEGATORS_ES if language == "es" else NEGATORS_EN

    pos, neg = 0, 0
    for i, tok in enumerate(tokens):
        # ventana de hasta 3 palabras previas para detectar negación
        window = tokens[max(0, i - 3):i]
        negated = any(w in negators for w in window)
        if tok in pos_w:
            neg += 1 if negated else 0
            pos += 0 if negated else 1
        elif tok in neg_w:
            pos += 1 if negated else 0
            neg += 0 if negated else 1

    if pos > neg:
        return "positivo" if language == "es" else "positive"
    if neg > pos:
        return "negativo" if language == "es" else "negative"
    return "neutro" if language == "es" else "neutral"

def sentiments_for_text(sentences: List[str], language: str) -> Dict[str, int]:
    keys = ("positivo", "negativo", "neutro") if language == "es" \
           else ("positive", "negative", "neutral")
    counts: Dict[str, int] = {k: 0 for k in keys}
    for s in sentences:
        counts[sentiment_score(s, language)] += 1
    return counts

def confidence_level(n: int, language: str) -> str:
    """3 niveles de confianza según cantidad de frases de evidencia."""
    if n >= 3:
        return "alta" if language == "es" else "high"
    if n >= 1:
        return "media" if language == "es" else "medium"
    return "baja" if language == "es" else "low"

def normalize_scores(scores: List[float]) -> List[float]:
    """Normalización min-max a [0,1] para poder comparar categorías de escalas distintas."""
    if not scores:
        return []
    lo, hi = min(scores), max(scores)
    if is_tie(hi, lo):
        return [0.5] * len(scores)
    return [(s - lo) / (hi - lo) for s in scores]

def get_colors(scores: List[float]) -> List[str]:
    if not scores:
        return []
    max_s, min_s = max(scores), min(scores)
    if is_tie(max_s, min_s):
        return ["grey"] * len(scores)
    result = []
    for s in scores:
        if is_tie(s, max_s):
            result.append("green")
        elif is_tie(s, min_s):
            result.append("red")
        else:
            result.append("orange")
    return result


# ═════════════════════════════════════════════
# Scoring numérico / cualitativo
# ═════════════════════════════════════════════
def score_category(sentences_list: List[List[str]], cat: str, language: str) -> List[float]:
    scores = []
    for sentences in sentences_list:
        total = 0.0
        if cat in NUMERIC_CATS:
            for s in sentences:
                nums = [float(n) for n in re.findall(r'\d+\.?\d*', s)]
                s_low = s.lower()
                if cat in ["pantalla", "screen"]:
                    total += nums[0] if nums else 0
                    total += sum(n / 1000 for n in nums if n > 1000)
                elif cat in ["cámara", "camera"]:
                    mp = re.findall(r'(\d+)\s*mp', s_low)
                    total += sum(float(n) for n in mp) if mp else 0
                elif cat in ["batería", "battery"]:
                    mah = re.findall(r'(\d{4,5})\s*mah', s_low)
                    watt = re.findall(r'(\d+)\s*w\b', s_low)
                    total += sum(float(n) for n in mah)
                    total += sum(float(n) for n in watt) * 10
                elif cat in ["rendimiento", "performance"]:
                    ram_nums = re.findall(r'(\d+)\s*gb\s*ram', s_low)
                    total += sum(float(n) for n in ram_nums)
                elif cat in ["almacenamiento", "storage"]:
                    storage_nums = re.findall(
                        r'(\d+)\s*(?:gb|tb)\s+(?:de\s+)?(?:almacenamiento|rom|memoria|storage|memory)', s_low)
                    total += sum(float(n) for n in storage_nums)
                elif cat in ["conectividad", "connectivity"]:
                    g = re.findall(r'\b(\d+)g\b', s_low)
                    wifi = re.findall(r'wifi\s*(\d+)', s_low)
                    bt = re.findall(r'bluetooth\s*(\d+\.?\d*)', s_low)
                    total += sum(float(n) for n in g + wifi + bt)
        else:
            # categorías cualitativas (precio, diseño, software, audio, durabilidad):
            # el score es el sentimiento neto (positivas - negativas) de las frases de evidencia
            for s in sentences:
                sent = sentiment_score(s, language)
                if sent in ("positivo", "positive"):
                    total += 1
                elif sent in ("negativo", "negative"):
                    total -= 1
        scores.append(total)
    return scores


# ═════════════════════════════════════════════
# Comparación principal
# ═════════════════════════════════════════════
def compare_by_categories(titles: List[str], texts: List[str], language: str = "es",
                           custom_categories: Optional[Dict[str, List[str]]] = None) -> Dict:
    texts = [clean_text(t) for t in texts]
    sents_list = [split_sentences(t) for t in texts]
    kws_dict = custom_categories if custom_categories else (CATEGORIES_ES if language == "es" else CATEGORIES_EN)

    cat_map_list = []
    for sents in sents_list:
        cat_map = {c: [] for c in kws_dict}
        for s in sents:
            for c in detect_categories(s, language, kws_dict):
                cat_map[c].append(s)
        cat_map_list.append(cat_map)

    results = {}
    for cat in kws_dict:
        sentences_per_product = [m.get(cat, []) for m in cat_map_list]
        confidence = [confidence_level(len(s), language) for s in sentences_per_product]
        scores = score_category(sentences_per_product, cat, language)
        colors = get_colors(scores)
        winners = [] if all(c == "grey" for c in colors) else \
            [titles[i] for i, c in enumerate(colors) if c == "green"]
        results[cat] = {
            "scores": scores,
            "normalized": normalize_scores(scores),
            "evidence": sentences_per_product,
            "confidence": confidence,
            "winners": winners,
        }

    victories = [0] * len(titles)
    for r in results.values():
        for w in r["winners"]:
            victories[titles.index(w)] += 1

    max_v = max(victories) if victories else 0
    winners_overall = [titles[i] for i, v in enumerate(victories) if v == max_v]
    overall = "Empate" if len(winners_overall) == len(titles) else ", ".join(winners_overall)

    return {
        "per_category": results,
        "overall": overall,
        "victories": victories,
        "categories": list(kws_dict.keys()),
    }

def weighted_ranking(titles: List[str], comp: Dict, weights: Dict[str, float]) -> List[Dict]:
    """Ranking ponderado usando los scores normalizados de cada categoría y pesos definidos por el usuario."""
    totals = [0.0] * len(titles)
    for cat, info in comp["per_category"].items():
        w = weights.get(cat, 1.0)
        for i, norm in enumerate(info["normalized"]):
            totals[i] += norm * w
    ranking = sorted(
        [{"title": titles[i], "score": round(totals[i], 3)} for i in range(len(titles))],
        key=lambda x: x["score"], reverse=True,
    )
    return ranking

def generate_summary(titles: List[str], comp: Dict, language: str) -> str:
    """Resumen automático en lenguaje natural a partir de qué producto gana en cada categoría."""
    won_map = {t: [] for t in titles}
    for cat, info in comp["per_category"].items():
        if len(info["winners"]) == 1:
            won_map[info["winners"][0]].append(cat)

    parts = []
    verbo = "gana en" if language == "es" else "wins in"
    for t in titles:
        cats = won_map[t]
        if cats:
            parts.append(f"**{t}** {verbo} {', '.join(cats)}")

    if not parts:
        return "No hay diferencias claras entre los productos." if language == "es" \
            else "No clear differences between the products."
    return " · ".join(parts)


# ═════════════════════════════════════════════
# Caché MD5 (con clave robusta vía JSON)
# ═════════════════════════════════════════════
def cached_compare(titles: List[str], texts: List[str], language: str,
                    custom_categories: Optional[Dict[str, List[str]]] = None) -> Dict:
    key_src = json.dumps(
        {"titles": titles, "texts": texts, "language": language, "custom_categories": custom_categories},
        sort_keys=True, ensure_ascii=False,
    )
    key = hashlib.md5(key_src.encode()).hexdigest()
    try:
        with shelve.open(CACHE_FILE) as db:
            if key in db:
                return db[key]
            result = compare_by_categories(titles, texts, language, custom_categories)
            db[key] = result
            return result
    except Exception:
        return compare_by_categories(titles, texts, language, custom_categories)

def clear_cache() -> str:
    removed = 0
    for f in glob.glob(CACHE_FILE + "*"):
        try:
            os.remove(f)
            removed += 1
        except Exception:
            pass
    return f"🗑️ Caché eliminada ({removed} archivo(s))."


# ═════════════════════════════════════════════
# Exportar CSV
# ═════════════════════════════════════════════
def export_csv(titles: List[str], comp: Dict) -> str:
    rows = []
    for cat, info in comp["per_category"].items():
        row = {"Categoría": cat}
        for i, title in enumerate(titles):
            row[f"Score {title}"] = round(info["scores"][i], 2)
            row[f"Normalizado {title}"] = round(info["normalized"][i], 2)
            row[f"Confianza {title}"] = info["confidence"][i]
        rows.append(row)
    df = pd.DataFrame(rows)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".csv", mode="w", encoding="utf-8")
    df.to_csv(tmp.name, index=False)
    tmp.close()
    return tmp.name


# ═════════════════════════════════════════════
# Gráficos (matplotlib, opcional)
# ═════════════════════════════════════════════
def make_radar_chart(titles: List[str], comp: Dict):
    if not HAS_MATPLOTLIB:
        return None
    cats = list(comp["per_category"].keys())
    if not cats:
        return None
    n = len(cats)
    angles = [i / n * 2 * 3.14159265 for i in range(n)]
    angles += angles[:1]

    fig = plt.figure(figsize=(6, 6))
    ax = fig.add_subplot(111, polar=True)
    for i, title in enumerate(titles):
        values = [comp["per_category"][c]["normalized"][i] for c in cats]
        values += values[:1]
        ax.plot(angles, values, label=title, linewidth=2)
        ax.fill(angles, values, alpha=0.1)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels([c.capitalize() for c in cats], fontsize=8)
    ax.set_yticklabels([])
    ax.legend(loc="upper right", bbox_to_anchor=(1.35, 1.1), fontsize=8)
    fig.tight_layout()
    return fig

def make_bar_chart(titles: List[str], comp: Dict):
    if not HAS_MATPLOTLIB:
        return None
    cats = list(comp["per_category"].keys())
    if not cats:
        return None
    fig, ax = plt.subplots(figsize=(7, max(3, len(cats) * 0.6)))
    y_pos = np.arange(len(cats))
    width = 0.8 / max(len(titles), 1)
    for i, title in enumerate(titles):
        scores = [comp["per_category"][c]["normalized"][i] for c in cats]
        ax.barh(y_pos + i * width, scores, height=width, label=title)
    ax.set_yticks(y_pos + width * (len(titles) - 1) / 2)
    ax.set_yticklabels([c.capitalize() for c in cats])
    ax.set_xlabel("Score normalizado")
    ax.invert_yaxis()
    ax.legend(fontsize=8)
    fig.tight_layout()
    return fig

def save_fig_png(fig) -> Optional[str]:
    if fig is None:
        return None
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    fig.savefig(tmp.name, dpi=130, bbox_inches="tight")
    tmp.close()
    return tmp.name


# ═════════════════════════════════════════════
# Exportar Word (.docx, opcional)
# ═════════════════════════════════════════════
def export_docx(titles: List[str], comp: Dict, language: str, radar_png: Optional[str] = None) -> Optional[str]:
    if not HAS_DOCX:
        return None
    doc = Document()
    doc.add_heading(
        f"{'Comparación' if language == 'es' else 'Comparison'}: {' vs '.join(titles)}", level=1
    )
    doc.add_paragraph(
        f"{'Resultado general' if language == 'es' else 'Overall result'}: {comp['overall']}"
    )
    if radar_png and os.path.exists(radar_png):
        doc.add_picture(radar_png, width=Inches(5.5))

    for cat, info in comp["per_category"].items():
        doc.add_heading(cat.capitalize(), level=2)
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "Producto" if language == "es" else "Product"
        hdr[1].text = "Score"
        hdr[2].text = "Confianza" if language == "es" else "Confidence"
        for i, title in enumerate(titles):
            row = table.add_row().cells
            row[0].text = title
            row[1].text = str(round(info["scores"][i], 2))
            row[2].text = info["confidence"][i]
        for i, evid in enumerate(info["evidence"]):
            if evid:
                doc.add_paragraph(f"{titles[i]}:")
                for s in evid:
                    doc.add_paragraph(f"• {s}", style="List Bullet")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


# ═════════════════════════════════════════════
# Lectura de archivos subidos (.txt / .pdf)
# ═════════════════════════════════════════════
def load_file_text(file):
    if file is None:
        return gr.update()
    path = file.name if hasattr(file, "name") else file
    try:
        if path.lower().endswith(".pdf"):
            if not HAS_PYPDF:
                return gr.update(value="[Error: instala 'pypdf' para leer PDFs]")
            reader = pypdf.PdfReader(path)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        else:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        return gr.update(value=text[:MAX_TEXT_CHARS])
    except Exception as e:
        return gr.update(value=f"[Error leyendo archivo: {e}]")


# ═════════════════════════════════════════════
# Construcción de outputs
# ═════════════════════════════════════════════
def build_outputs(titles: List[str], texts: List[str], language: str,
                   custom_categories: Optional[Dict[str, List[str]]] = None,
                   weights: Optional[Dict[str, float]] = None):
    comp = cached_compare(titles, texts, language, custom_categories)
    csv_path = export_csv(titles, comp)
    weights = weights or {}
    ranking = weighted_ranking(titles, comp, weights)
    summary = generate_summary(titles, comp, language)

    radar_fig = make_radar_chart(titles, comp)
    bar_fig = make_bar_chart(titles, comp)
    radar_png = save_fig_png(radar_fig)
    docx_path = export_docx(titles, comp, language, radar_png)

    md = [
        f"# {'Comparación' if language == 'es' else 'Comparison'}: {' vs '.join(titles)}",
        f"**{'Resultado general' if language == 'es' else 'Overall result'}:** {comp['overall']}",
        "",
        f"### {'Resumen automático' if language == 'es' else 'Automatic summary'}",
        summary,
        "",
        f"### {'Ranking ponderado' if language == 'es' else 'Weighted ranking'}",
    ]
    for i, r in enumerate(ranking, start=1):
        md.append(f"{i}. **{r['title']}** — {r['score']}")

    md += ["", f"### {'Sentimiento por texto' if language == 'es' else 'Sentiment per text'}"]
    for title, text in zip(titles, texts):
        sents = split_sentences(text)
        counts = sentiments_for_text(sents, language)
        md.append(f"**{title}**: " + " · ".join(f"{k}: {v}" for k, v in counts.items()))

    md += ["", f"### {'Especificaciones detectadas' if language == 'es' else 'Detected specs'}"]
    for title, text in zip(titles, texts):
        specs = extract_specs(text, language)
        if specs:
            agg = aggregate_specs(specs)
            parts = []
            for k, v in agg.items():
                extra = f" ({', '.join(str(x) for x in specs[k])})" if len(specs[k]) > 1 else ""
                parts.append(f"{k}={v:g}{extra}")
            line = ", ".join(parts)
        else:
            line = "—"
        md.append(f"**{title}**: {line}")
    md.append("---")

    conf_color = {"alta": "green", "media": "orange", "baja": "red",
                  "high": "green", "medium": "orange", "low": "red"}

    html_rows = ["<tr><th>Categoría</th><th>Ganador</th><th>Confianza</th></tr>"]
    for cat, info in comp["per_category"].items():
        colors = get_colors(info["scores"])
        winner_text = "Empate" if not info["winners"] else ", ".join(info["winners"])
        conf_text = " / ".join(
            f"<span style='color:{conf_color.get(c, 'gray')}'>{c}</span>" for c in info["confidence"]
        )
        md.append(f"## {cat.capitalize()}")
        md.append(f"**{'Ganador' if language == 'es' else 'Winner'}:** {winner_text}")
        for i, evid in enumerate(info["evidence"]):
            if evid:
                md.append(f"**{titles[i]} — Evidencia:**")
                for s in evid:
                    sent = sentiment_score(s, language)
                    sent_color = (
                        "green" if sent in ("positivo", "positive") else
                        "red" if sent in ("negativo", "negative") else
                        "gray"
                    )
                    md.append(
                        f"- <span style='color:{colors[i]}'>{s}</span> "
                        f"<small style='color:{sent_color}'>[{sent}]</small>"
                    )
        html_rows.append(f"<tr><td>{cat}</td><td>{winner_text}</td><td>{conf_text}</td></tr>")
        md.append("")

    html = ("<table style='width:100%;border-collapse:collapse;font-size:14px;'>"
            + "".join(html_rows) + "</table>")
    json_out = json.dumps(comp, ensure_ascii=False, indent=2)

    return md_join(md), html, json_out, csv_path, radar_fig, bar_fig, docx_path

def md_join(md_lines: List[str]) -> str:
    return "\n".join(md_lines)


# ═════════════════════════════════════════════
# Historial de sesión
# ═════════════════════════════════════════════
def render_history(history: List[Dict]) -> str:
    if not history:
        return "_Sin comparaciones previas en esta sesión._"
    lines = ["| Hora | Productos | Resultado |", "|---|---|---|"]
    for h in reversed(history[-20:]):
        lines.append(f"| {h['timestamp']} | {h['titles']} | {h['overall']} |")
    return "\n".join(lines)


# ═════════════════════════════════════════════
# Gradio UI
# ═════════════════════════════════════════════
def run_gradio(product_count, lang_override, custom_cats_text, weights_text, history, *product_args):
    titles, texts = [], []
    for i in range(int(product_count)):
        t = product_args[i * 2]
        x = product_args[i * 2 + 1]
        if x and x.strip():
            xx = x if len(x) <= MAX_TEXT_CHARS else x[:MAX_TEXT_CHARS]
            titles.append(t or f"Producto {chr(65 + i)}")
            texts.append(xx)

    history = list(history or [])

    if len(texts) < 2:
        return ("⚠️ Introduce texto en al menos 2 productos.", "", "{}", None,
                None, None, None, history, render_history(history))

    language = detect_language(texts) if lang_override == "auto" else lang_override

    custom_categories = None
    if custom_cats_text and custom_cats_text.strip():
        try:
            custom_categories = json.loads(custom_cats_text)
        except Exception:
            custom_categories = None  # JSON inválido -> se ignora y se usan categorías por defecto

    weights = {}
    if weights_text and weights_text.strip():
        try:
            weights = json.loads(weights_text)
        except Exception:
            weights = {}

    md, html, json_out, csv_path, radar_fig, bar_fig, docx_path = build_outputs(
        titles, texts, language, custom_categories, weights
    )

    comp_overall = json.loads(json_out)["overall"]
    history.append({
        "timestamp": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S"),
        "titles": " vs ".join(titles),
        "overall": comp_overall,
    })

    return md, html, json_out, csv_path, radar_fig, bar_fig, docx_path, history, render_history(history)


def update_rows(n):
    return [gr.update(visible=(i < int(n))) for i in range(MAX_PRODUCTS)]


with gr.Blocks(title="Comparador Inteligente", theme=gr.themes.Soft()) as demo:
    gr.Markdown("# Comparador Inteligente\nCompara hasta 5 productos por categorías")

    with gr.Row():
        product_count = gr.Slider(2, MAX_PRODUCTS, value=2, step=1,
                                   label="Número de productos")
        lang_radio = gr.Radio(["auto", "es", "en"], value="auto", label="Idioma",
                               info="auto = detección automática con langdetect")

    with gr.Accordion("Opciones avanzadas", open=False):
        custom_cats_box = gr.Textbox(
            label="Categorías personalizadas (JSON, opcional)",
            placeholder='{"cámara": ["cámara", "mp", "sensor"], "batería": ["batería", "mah"]}',
            lines=3,
        )
        weights_box = gr.Textbox(
            label="Pesos por categoría para el ranking ponderado (JSON, opcional; por defecto 1.0)",
            placeholder='{"cámara": 2, "batería": 1.5, "pantalla": 1}',
            lines=2,
        )
        with gr.Row():
            clear_cache_btn = gr.Button("🗑️ Limpiar caché")
            clear_cache_status = gr.Markdown()
        clear_cache_btn.click(fn=clear_cache, outputs=clear_cache_status)

    product_inputs = []
    product_rows = []
    for i in range(MAX_PRODUCTS):
        with gr.Row(visible=(i < 2)) as row:
            t_box = gr.Textbox(label=f"Título {chr(65 + i)}",
                                value=f"Producto {chr(65 + i)}", scale=1)
            x_box = gr.Textbox(label=f"Texto {chr(65 + i)}", lines=6, scale=3)
            f_box = gr.File(label=f"Archivo {chr(65 + i)} (.txt/.pdf)", scale=1,
                             file_types=[".txt", ".pdf"])
        f_box.upload(fn=load_file_text, inputs=f_box, outputs=x_box)
        product_rows.append(row)
        product_inputs.extend([t_box, x_box])

    product_count.change(fn=update_rows, inputs=product_count, outputs=product_rows)

    btn = gr.Button("Comparar", variant="primary")

    history_state = gr.State([])

    with gr.Tabs():
        with gr.Tab("Resumen"):
            md_out = gr.Markdown()
        with gr.Tab("Tabla"):
            html_out = gr.HTML()
        with gr.Tab("Gráficos"):
            radar_plot = gr.Plot(label="Radar por categoría (normalizado)")
            bar_plot = gr.Plot(label="Barras por categoría (normalizado)")
        with gr.Tab("JSON"):
            json_out = gr.Textbox(label="JSON completo", lines=20)
        with gr.Tab("Exportar"):
            csv_out = gr.File(label="Descargar CSV")
            docx_out = gr.File(label="Descargar informe Word (.docx)")
        with gr.Tab("Historial"):
            history_md = gr.Markdown(render_history([]))

    btn.click(
        fn=run_gradio,
        inputs=[product_count, lang_radio, custom_cats_box, weights_box, history_state] + product_inputs,
        outputs=[md_out, html_out, json_out, csv_out, radar_plot, bar_plot, docx_out, history_state, history_md],
    )

if __name__ == "__main__":
    demo.launch()
